"""
File Manager Module for Finance Accountant Agent

This module handles the ingestion and processing of various file types for RAG.
It supports PDF, Excel, Word, and CSV files, extracting text and metadata for
efficient retrieval.

Features:
- Multiple file format support (PDF, Excel, Word, CSV)
- Metadata extraction from files
- Text chunking with appropriate overlap
- Integration with RAG module for embedding
- Batch processing for multiple files
- Error handling for various file issues

Dependencies:
- rag_module: For adding processed documents to vector store
- PyPDF2/pdfplumber: For PDF processing
- openpyxl: For Excel processing
- python-docx: For Word document processing
- pandas: For CSV processing
"""

import asyncio
import csv
import io
import logging
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import pandas as pd
from fastapi import UploadFile

from config.settings import settings
from modules.rag_module import Document, rag_module

logger = logging.getLogger(__name__)


async def process_pdf(file_content: bytes, filename: str, metadata: Dict) -> List[Document]:
    try:
        import pdfplumber
        documents = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    doc_metadata = {
                        **metadata,
                        "source": filename,
                        "page": i + 1,
                        "total_pages": len(pdf.pages),
                    }
                    if len(text) > settings.RAG_CHUNK_SIZE:
                        chunks = split_text(text, settings.RAG_CHUNK_SIZE, settings.RAG_CHUNK_OVERLAP)
                        for j, chunk in enumerate(chunks):
                            chunk_metadata = {**doc_metadata, "chunk": j + 1, "total_chunks": len(chunks)}
                            documents.append(Document(text=chunk, metadata=chunk_metadata))
                    else:
                        documents.append(Document(text=text, metadata=doc_metadata))
        return documents
    except Exception as e:
        logger.error(f"Error processing PDF file {filename}: {str(e)}")
        raise


async def process_excel(file_content: bytes, filename: str, metadata: Dict) -> List[Document]:
    try:
        documents = []
        xlsx = pd.ExcelFile(io.BytesIO(file_content))
        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(xlsx, sheet_name=sheet_name)
            if df.empty:
                continue
            text = f"Sheet: {sheet_name}\n\n" + df.to_string(index=False)
            doc_metadata = {
                **metadata,
                "source": filename,
                "sheet": sheet_name,
            }
            if len(text) > settings.RAG_CHUNK_SIZE:
                chunks = split_text(text, settings.RAG_CHUNK_SIZE, settings.RAG_CHUNK_OVERLAP)
                for j, chunk in enumerate(chunks):
                    chunk_metadata = {**doc_metadata, "chunk": j + 1, "total_chunks": len(chunks)}
                    documents.append(Document(text=chunk, metadata=chunk_metadata))
            else:
                documents.append(Document(text=text, metadata=doc_metadata))
        return documents
    except Exception as e:
        logger.error(f"Error processing Excel file {filename}: {str(e)}")
        raise


async def process_word(file_content: bytes, filename: str, metadata: Dict) -> List[Document]:
    try:
        import docx
        documents = []
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        try:
            doc = docx.Document(temp_file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text])
            doc_metadata = {
                **metadata,
                "source": filename,
            }
            if len(full_text) > settings.RAG_CHUNK_SIZE:
                chunks = split_text(full_text, settings.RAG_CHUNK_SIZE, settings.RAG_CHUNK_OVERLAP)
                for j, chunk in enumerate(chunks):
                    chunk_metadata = {**doc_metadata, "chunk": j + 1, "total_chunks": len(chunks)}
                    documents.append(Document(text=chunk, metadata=chunk_metadata))
            else:
                documents.append(Document(text=full_text, metadata=doc_metadata))
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
        return documents
    except Exception as e:
        logger.error(f"Error processing Word file {filename}: {str(e)}")
        raise


async def process_csv(file_content: bytes, filename: str, metadata: Dict) -> List[Document]:
    try:
        documents = []
        df = pd.read_csv(io.BytesIO(file_content))
        if df.empty:
            return []
        text = df.to_string(index=False)
        doc_metadata = {
            **metadata,
            "source": filename,
        }
        if len(text) > settings.RAG_CHUNK_SIZE:
            chunks = split_text(text, settings.RAG_CHUNK_SIZE, settings.RAG_CHUNK_OVERLAP)
            for j, chunk in enumerate(chunks):
                chunk_metadata = {**doc_metadata, "chunk": j + 1, "total_chunks": len(chunks)}
                documents.append(Document(text=chunk, metadata=chunk_metadata))
        else:
            documents.append(Document(text=text, metadata=doc_metadata))
        return documents
    except Exception as e:
        logger.error(f"Error processing CSV file {filename}: {str(e)}")
        raise


def split_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if end < len(text):
            break_chars = ["\n\n", "\n", ".", "!", "?", ",", " "]
            for break_char in break_chars:
                last_break = chunk.rfind(break_char)
                if last_break != -1:
                    end = start + last_break + 1
                    chunk = text[start:end]
                    break
        chunks.append(chunk)
        start = end - chunk_overlap
    return chunks


async def ingest_file(file: UploadFile, category: str) -> Dict:
    try:
        content = await file.read()
        file_size_mb = len(content) / (1024 * 1024)
        if file_size_mb > settings.MAX_UPLOAD_SIZE_MB:
            raise Exception(f"File too large ({file_size_mb:.1f} MB). Maximum allowed size is {settings.MAX_UPLOAD_SIZE_MB} MB")

        await file.seek(0)
        filename = file.filename or "unknown_file"
        file_extension = os.path.splitext(filename)[1].lower()

        if file_extension not in settings.ALLOWED_FILE_TYPES:
            raise Exception(f"Unsupported file type: {file_extension}. Allowed types: {', '.join(settings.ALLOWED_FILE_TYPES)}")

        metadata = {
            "category": category,
            "date": datetime.now().isoformat(),
            "filename": filename,
            "file_type": file_extension,
            "file_size": len(content),
        }

        documents = []

        if file_extension == ".pdf":
            documents = await process_pdf(content, filename, metadata)
        elif file_extension in [".xlsx", ".xls"]:
            documents = await process_excel(content, filename, metadata)
        elif file_extension == ".docx":
            documents = await process_word(content, filename, metadata)
        elif file_extension == ".csv":
            documents = await process_csv(content, filename, metadata)
        elif file_extension == ".txt":
            text = content.decode("utf-8", errors="replace")
            chunks = split_text(text, settings.RAG_CHUNK_SIZE, settings.RAG_CHUNK_OVERLAP)
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **metadata,
                    "chunk": i + 1,
                    "total_chunks": len(chunks),
                }
                documents.append(Document(text=chunk, metadata=chunk_metadata))

        num_added = await rag_module.add_documents(documents)

        return {
            "filename": filename,
            "file_type": file_extension,
            "category": category,
            "chunks_added": num_added,
            "status": "success",
        }

    except Exception as e:
        logger.error(f"Error ingesting file {file.filename}: {str(e)}")
        raise
